"""Generate a per-mood clinical-review Word document for TherapiShots."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

MOODS = [
    {"key": "heavy", "emoji": "😔", "group": "low", "value": 1, "en": "Heavy"},
    {"key": "anxious", "emoji": "😰", "group": "low", "value": 2, "en": "Anxious"},
    {"key": "frustrated", "emoji": "😤", "group": "low", "value": 2, "en": "Frustrated"},
    {"key": "numb", "emoji": "😶", "group": "neutral", "value": 3, "en": "Numb"},
    {"key": "foggy", "emoji": "😑", "group": "neutral", "value": 3, "en": "Foggy"},
    {"key": "okay", "emoji": "😐", "group": "neutral", "value": 4, "en": "Okay"},
    {"key": "hopeful", "emoji": "🌱", "group": "bright", "value": 5, "en": "Hopeful"},
    {"key": "calm", "emoji": "😌", "group": "bright", "value": 5, "en": "Calm"},
    {"key": "energised", "emoji": "⚡", "group": "bright", "value": 6, "en": "Energised"},
]

GROUP_DESC = {
    "low": "Lower-mood / distress-leaning state",
    "neutral": "Flat, in-between or unclear state",
    "bright": "Positive / resourced state",
}


def small_steps_for(value):
    if value <= 2:
        return "2-minute breathing  OR  Contact someone you trust"
    if value == 3:
        return "Short walk  OR  Step outside"
    return "Stretch  /  Hydrate  /  Step outside"


def day_notice_for(group):
    if group == "low":
        return ("• Single low check-in today → \u201cToday landed on the heavier side. "
                "Thank you for noticing it.\u201d\n"
                "• \u2265 2 low check-ins today → repeated-low nudge; if a low mood also "
                "appears on \u22652 of the last 4 days → suggests talking to someone.\n"
                "• If a bright mood is also logged the same day → \u201cemotional variability\u201d "
                "message (normalises ups & downs).")
    if group == "neutral":
        return ("• A day of only neutral/foggy check-ins → \u201cA flat or foggy stretch today. "
                "Gentle routine, rest, and a little daylight can help the edges feel sharper.\u201d\n"
                "• If mixed with bright (no low) → \u201cmostly steady day\u201d message.")
    return ("• A day of only bright check-ins → \u201cToday has felt steadily brighter. Worth "
            "pausing to notice what\u2019s working.\u201d\n"
            "• If a low mood also appears same day → \u201cemotional variability\u201d message.")


def call_trigger_for(key, group):
    if group == "low":
        return ("YES — selecting this mood immediately shows the inline support prompt and "
                "recommends a 15-minute psychologist call (this mood is in the "
                "low-mood set: Heavy, Anxious, Frustrated). Also contributes to the "
                "repeated-low signal (\u22653 days below personal baseline\u22120.5 within the last 7).")
    return ("Not immediately. Does not itself trigger the call prompt, but still counts "
            "toward the personal baseline used by repeated-low detection.")


def add_heading_color(doc, text, size, rgb):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*rgb)
    return p


doc = Document()

# Title
add_heading_color(doc, "TherapiShots — Mood-by-Mood Reflection Logic", 20, (0x2C, 0x24, 0x16))
sub = doc.add_paragraph("Clinical review document — how each mood is defined and which reflections, "
                        "prompts and support actions it triggers in the app.")
sub.runs[0].italic = True
doc.add_paragraph("Prepared for review with a Clinical Psychologist.   Version 1.0")
doc.add_paragraph(
    "Notes: (1) All reflections are deterministic (rule-based), not AI-generated, and are "
    "explicitly non-diagnostic. (2) \u201cValue\u201d is an internal 0\u20136 scale (higher = feeling better) "
    "used only for statistics; it is never shown as a score to the user. (3) Health signals are "
    "simulated demo data and are used only in the Wellbeing Pulse, never in the pattern insights.")
doc.add_paragraph("Mood groups:  Low (Heavy, Anxious, Frustrated) · Neutral (Numb, Foggy, Okay) · "
                  "Bright (Hopeful, Calm, Energised).")

doc.add_page_break()

for m in MOODS:
    add_heading_color(doc, f"{m['emoji']}  {m['en']}", 16, (0x3D, 0x4F, 0x7C))
    meta = doc.add_paragraph()
    meta.add_run("Group: ").bold = True
    meta.add_run(f"{m['group'].capitalize()} — {GROUP_DESC[m['group']]}    ")
    meta.add_run("Internal value: ").bold = True
    meta.add_run(f"{m['value']}/6")

    rows = [
        ("Immediate support / 15-min call trigger", call_trigger_for(m['key'], m['group'])),
        ("\u201cSmall step\u201d suggested", small_steps_for(m['value'])),
        ("Feel Map (Progress)", f"Adds a {m['group']}-coloured dot for that day (last 42 days)."),
        ("Mood trend (Progress)", f"Contributes value {m['value']}/6 to the 30-day bar chart "
                                  f"(taller = brighter)."),
        ("Insights engine (from 7+ check-ins)", "Feeds day-of-week and context-tag patterns "
                                                "(e.g. a context tag chosen with this mood \u22653 times "
                                                "may surface as \u2018appears in lower-mood check-ins\u2019)."),
        ("Story composition (Read your story)", f"Counted as a \u2018{m['group']}\u2019 day; \u226550% bright "
                                                f"\u2192 bright, \u226540% low \u2192 low, \u226550% neutral \u2192 steady, "
                                                f"else mixed."),
        ("Day-notice (\u2018Something you may want to notice\u2019)", day_notice_for(m['group'])),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, val in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(2.2)
        cells[1].width = Inches(4.3)
        rp = cells[0].paragraphs[0].add_run(label)
        rp.bold = True
        cells[1].paragraphs[0].add_run(val)

    cn = doc.add_paragraph()
    cn.add_run("Clinical reviewer notes: ").bold = True
    doc.add_paragraph("________________________________________________________________")
    doc.add_paragraph("________________________________________________________________")
    doc.add_paragraph("")

# Appendix
doc.add_page_break()
add_heading_color(doc, "Appendix — Trigger thresholds (by number of check-ins, n)", 15, (0x2C, 0x24, 0x16))
appendix = [
    ("Save confirmation + small step", "n \u2265 1 (every check-in)"),
    ("Day-notice", "\u2265 1 check-in today (logic by mood mix; escalates on \u22652 low of last 4 days)"),
    ("Low-mood prompt / 15-min call", "Latest mood \u2208 {Heavy, Anxious, Frustrated}  OR  repeated-low "
                                      "(\u22655 check-ins & \u22653 days below baseline\u22120.5 in last 7)"),
    ("Today observation", "n \u2265 5"),
    ("Insights engine ON", "n \u2265 7"),
    ("Context pattern", "a tag seen \u2265 3\u00d7 (Emerging at \u2265 6)"),
    ("Weekday pattern", "n \u2265 10 (Emerging at \u2265 15)"),
    ("Confidence badges", "Early Signal 7\u201314 · Emerging 15\u201329 · Consistent \u2265 30 (|effect| \u2265 0.18)"),
    ("Feel Map / Mood trend", "Any history (42-day map · 30-day trend)"),
    ("Story trend", "\u2265 4 mood values in the window (week = 7 days, month = 30 days)"),
    ("AI rephrase of story", "Only if \u2018AI summaries\u2019 consent is ON (rephrases validated facts only)"),
]
t2 = doc.add_table(rows=0, cols=2)
t2.style = "Light Grid Accent 1"
for a, b in appendix:
    c = t2.add_row().cells
    c[0].paragraphs[0].add_run(a).bold = True
    c[1].paragraphs[0].add_run(b)

out = "/app/TherapiShots_Mood_Clinical_Review.docx"
doc.save(out)
print("saved", out)
